"""Generate AlphaAgent_Pitch_Script.docx: slide-by-slide speaker script + glossary."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEAL = RGBColor(0x0F, 0x8A, 0x7D)   # darker teal for print legibility
NAVY = RGBColor(0x0E, 0x22, 0x3B)
GRAY = RGBColor(0x64, 0x74, 0x8B)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = NAVY
    return p

def slide_head(num, title):
    p = doc.add_paragraph(); p.space_before = Pt(14)
    r = p.add_run(f"SLIDE {num} — {title}")
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = TEAL

def script(text, time=None):
    if time:
        p = doc.add_paragraph()
        r = p.add_run(f"[{time}]"); r.font.size = Pt(9); r.font.color.rgb = GRAY; r.font.italic = True
    for para in text.strip().split("\n\n"):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.space_after = Pt(6)

# ---------------- Cover page ----------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("\n\n\nAlphaAgent Healthcare Solutions"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Investor Pitch — Presenter Script & Glossary"); r.font.size = Pt(14); r.font.color.rgb = TEAL
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("Companion to AlphaAgent_Pitch_Deck_v2.pptx  ·  Confidential  ·  July 2026")
r.font.size = Pt(10); r.font.color.rgb = GRAY
doc.add_page_break()

# ---------------- Script ----------------
h1("Presenter Script")
p = doc.add_paragraph()
r = p.add_run("Target: 10–12 minutes for the full walk-through; slides 1–3 in the first 90 seconds. "
              "Timing cues are cumulative. Speak it, don't read it — bullets on screen carry the detail.")
r.font.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY

slide_head(1, "TITLE")
script("""
I'm Nathanael Gill, CEO of AlphaAgent Healthcare Solutions. We build the autonomous documentation layer for the operating room: smart glasses plus an agentic AI that documents surgery completely hands-free. I'll show you why surgeons desperately want this, why it's only become possible in the last eighteen months, and why a small check now buys a position in a regulated market with real moats.
""", "0:00")

slide_head(2, "THE PROBLEM")
script("""
Surgeons today are documentation clerks with scalpels. Charting is the number-one driver of clinician burnout — hours every day lost to a keyboard. And splitting attention between the patient and the chart measurably increases error rates.

There's a second failure nobody talks about: pre- and post-op wound photos are routinely taken on personal phones — a personal device inside a sterile field. That's a contamination risk and a compliance nightmare.

Today's AI scribes stop at a draft note. A human still does all the clicking. The keyboard and the phone are both in the wrong place — we remove both.
""", "0:30")

slide_head(3, "THE SOLUTION")
script("""
Three steps: capture, reason, act. The surgeon wears Vuzix LX1 glasses — twelve-megapixel camera, OR-grade microphones, fully sterile because it's hands-free. A cloud agent with persistent memory turns that capture into a structured clinical record. Then — and this is the difference — the agent acts: it writes the EHR fields itself through the FHIR standard. The surgeon reviews and says "approve." Zero phones, zero keyboards. The chart finishes itself.
""", "1:30")

slide_head(4, "PRODUCT & TECHNICAL MOAT")
script("""
The glasses are a full Android 15 computer — eight-core processor, on-device pre-processing, so protected health information is handled before anything leaves the room. The agentic core is our own stack: persistent memory, deterministic tool-calling, FHIR writes — not a wrapper around a chat model. The infrastructure is compliance-native from day one: every vendor under a signed BAA. And every record a surgeon approves becomes labeled training data — a data flywheel nobody else in this market has.
""", "2:30")

slide_head(5, "WHY NOW")
script("""
Three curves crossed at the same time. AI agents crossed from chatting to reliably acting in 2025 and 2026. Enterprise eyewear finally has the compute, optics, and microphones for an OR. And hospitals are desperate — burnout and staffing shortages have buyers actively searching for exactly this. This product was technically impossible three years ago. The window is open, and it will not stay uncontested.
""", "3:30")

slide_head(6, "COMPETITIVE LANDSCAPE")
script("""
Map the market on two axes: hands-free capture, and agentic action. DAX and Abridge are voice scribes — no eyewear, no visual capture, and they stop at a draft note. Vuzix makes the best enterprise glasses but has no clinical agent. We're alone in the top-right quadrant: wearable AND agentic. For a scribe company to catch us they must become a hardware company; for a hardware company, they must build our agent. Either chase takes years — years we'll spend banking regulatory clearances.
""", "4:15")

slide_head(7, "MARKET OPPORTUNITY")
script("""
Healthcare is an eight-trillion-dollar global market; our serviceable segment is US clinical documentation and surgical workflow. Our beachhead — the SOM — is high-volume surgical specialties: orthopedics, general surgery, plastics, across US hospitals and ambulatory surgery centers. The wedge is owning the moment of capture in the OR. The platform is what that capture builds: a proprietary multimodal surgical dataset that powers automated diagnostics later.
""", "5:00")

slide_head(8, "BUSINESS MODEL")
script("""
B2B SaaS. The hospital or surgery center pays a site license of five to twenty thousand dollars a month depending on size — the institution pays, never the individual surgeon. The glasses are free to the hospital: at roughly two thousand dollars a unit and ten units per site, we recover the hardware cost in just over two months of subscription gross profit. Then it's pure software margin. And it's sticky by design — ripping us out means going back to the keyboard and the phone.
""", "5:45")

slide_head(9, "REGULATORY & PATH TO MARKET")
script("""
We treat regulation as a moat, not a tax. Pathway is a Class II 510(k) with surgical lighting and headwear as predicates. Lean ISO 13485 quality system from day one; HIPAA and SOC 2 posture built in, not bolted on. Q3 this year we file the FDA Pre-Submission, which locks the testing protocol. Honest milestone: this round funds a cleared pathway plus live pilots — full 510(k) clearance completes around 2027 on the next round. Every one of those steps is a barrier behind us.
""", "6:45")

slide_head(10, "TEAM")
script("""
A hardware-and-software product needs a hardware-and-software team. I build the agent — memory, tool-calling, FHIR integration — and I did healthcare analytics at HCA, so I know clinical data pipelines and HIPAA from the inside. Manny Figueroa, our CTO, is a mechanical and systems engineer from Raytheon — reliability-critical hardware — and he leads the custom enclosure phase. He's equity-only until the pilot proves the wedge. One salary, off-the-shelf hardware, custom build only after the pilot earns it.
""", "7:30")

slide_head(11, "FINANCIAL PROJECTIONS")
script("""
Bottom-up, and deliberately honest. Blended site license of ten thousand a month. Pilots take eight months — that's real hospital procurement and IT review, not optimism — and convert to paid in month nine. Three sites by end of year one, ten by end of year two, twenty-five by end of year three as the 510(k) unlocks the ramp. That's three million dollars of ARR exiting year three — and here's the line that matters: net income turns positive in year three. Total outside capital to breakeven is about two million dollars: this five-hundred-K pre-seed plus a roughly one-and-a-half-million seed raised on converted pilots. Most medtech companies burn ten million before clearance. We reach self-funding on two.
""", "8:15")

slide_head(12, "THE LONG VIEW")
script("""
Where does this go? There are roughly six thousand hospitals and eleven thousand ambulatory surgery centers in the US. Seven hundred fifty sites in ten years is four-point-four percent penetration — no heroics required. Contract value compounds inside each site as we expand from one service line to the department to the system standard of care: one-thirty-five million ARR by year ten. And that's before the second act — the proprietary surgical dataset and diagnostics licensing that stack on top of the subscription base. Sticky product, regulatory barriers, self-funding core, data moat. That's how this compounds.
""", "9:30")

slide_head(13, "THE ASK")
script("""
We're raising about five hundred thousand dollars. Roughly a hundred goes to the hardware fleet and tooling, two hundred to regulatory and compliance — quality system, HIPAA, SOC 2, 510(k) started — and two-twenty covers eighteen months of runway including Manny post-pilot and cloud costs. Three hundred K is the efficient floor: twelve months, FDA Pre-Sub, single-site pilots. The line I'll leave you with: we prove the wedge on a hundred K; we clear the regulatory moat on five hundred.
""", "10:30")

slide_head(14, "CLOSE")
script("""
Own the surgeon's point of view, and you own the future of the operating room. Today we kill documentation friction. Tomorrow, the intelligent capture-and-action layer for all of surgery. And always: building the surgical dataset that powers automated diagnostics. Thank you — happy to go deep on the model, the regulatory path, or the tech.
""", "11:15")

doc.add_page_break()

# ---------------- Glossary ----------------
h1("Glossary of Abbreviations & Terms")
terms = [
    ("ACV", "Annual Contract Value — total subscription revenue one customer (site) pays per year. A $10K/month site license = $120K ACV."),
    ("ARR", "Annual Recurring Revenue — all active subscriptions annualized; the core SaaS growth metric. 25 sites × $120K = $3.0M ARR."),
    ("ASC", "Ambulatory Surgery Center — outpatient facility for same-day surgery; ~11,000 in the US. Smaller, faster-moving buyers than hospitals."),
    ("B2B SaaS", "Business-to-Business Software as a Service — selling subscription software to institutions (hospitals) rather than consumers."),
    ("BAA", "Business Associate Agreement — HIPAA-required contract making a vendor (e.g., Azure, Google Cloud) legally liable for protecting patient data it touches."),
    ("Class II / 510(k)", "FDA clearance pathway for moderate-risk medical devices. A 510(k) shows the device is 'substantially equivalent' to an existing cleared device (the predicate). Months and hundreds of $K, vs. years and $M for the stricter PMA pathway."),
    ("COGS", "Cost of Goods Sold — direct cost of serving a customer: cloud/inference/support ($1K/site/mo) plus the free glasses ($20K per new site)."),
    ("EHR", "Electronic Health Record — the hospital's patient-record system (Epic, Cerner/Oracle Health). Our agent writes into it directly."),
    ("EOY", "End Of Year — figures measured at the last day of the year (e.g., sites live at year-end)."),
    ("FDA Pre-Sub", "Pre-Submission — a formal meeting where the FDA agrees in advance what testing your 510(k) must include. De-risks the whole regulatory path."),
    ("FHIR", "Fast Healthcare Interoperability Resources ('fire') — the modern API standard for reading/writing EHR data. It's how the agent files the chart."),
    ("GM / Gross margin", "Gross Margin — (revenue − COGS) ÷ revenue. Ours starts low (free hardware expensed up front) and climbs to ~90% on software."),
    ("HIPAA", "Health Insurance Portability and Accountability Act — US law governing patient-data privacy and security. Non-negotiable for hospital sales."),
    ("ISO 13485", "International quality-management standard for medical-device companies; the FDA expects an equivalent quality system (QMS)."),
    ("LX1", "Vuzix LX1 — the enterprise smart glasses we deploy: Android 15, 8-core CPU, 12MP camera, Wi-Fi 6E, ~$2K/unit."),
    ("Net income", "Revenue minus all costs (COGS + opex). Negative = burning cash ('burn'); positive = cashflow-positive/self-funding."),
    ("Opex", "Operating Expenses — everything outside COGS: salaries, regulatory fees, sales, G&A (general & administrative)."),
    ("OR", "Operating Room."),
    ("PHI", "Protected Health Information — patient-identifiable data (photos, audio, records). Everything HIPAA protects; we pre-process it on-device."),
    ("Pre-seed / Seed / Series A", "Successive startup funding rounds. Pre-seed (~$500K, this raise) proves the wedge; seed (~$1.5M) scales pilots to revenue; Series A ($5M+) is optional growth capital once we're self-funding."),
    ("QMS", "Quality Management System — the documented processes (design controls, audits, traceability) a medical-device company must run; ours follows ISO 13485."),
    ("Runway", "How many months the company can operate before cash runs out at the current burn rate."),
    ("SOC 2", "Service Organization Control 2 — an independent security audit hospital IT departments demand from software vendors."),
    ("SOM / SAM / TAM", "Serviceable Obtainable / Serviceable Addressable / Total Addressable Market — beachhead we can win now (US high-volume surgical specialties) ⊂ segment we can serve (US clinical documentation) ⊂ whole market ($8T global healthcare)."),
    ("US penetration", "Share of the ~17,000 US surgical facilities (hospitals + ASCs) we've signed. 750 sites in Year 10 = 4.4%."),
    ("Wedge", "The narrow first product that gets you into the customer (OR documentation), from which you expand to the platform (surgical intelligence)."),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for k, txt_ in ((0, "Term"), (1, "Meaning")):
    r = hdr[k].paragraphs[0].add_run(txt_); r.font.bold = True
tbl.columns[0].width = Inches(1.6); tbl.columns[1].width = Inches(5.2)
for term, meaning in terms:
    row = tbl.add_row().cells
    r = row[0].paragraphs[0].add_run(term); r.font.bold = True
    row[1].paragraphs[0].add_run(meaning)
    row[0].width = Inches(1.6); row[1].width = Inches(5.2)

doc.save("/Users/nathanaelguitar/Downloads/AlphaAgent_Pitch_Script.docx")
print("saved")
